import os
from typing import List, Dict, Any, Optional
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

class DocxExporter:
    """
    Exportador principal de documentos editables a formato Microsoft Word (.docx).
    Reconstruye la estructura del documento impreso original e inserta campos
    diligenciables (texto, fecha, firma, casillas de verificación).
    """

    def __init__(self):
        pass

    def create_base_document(self) -> docx.Document:
        """Crea un documento Word con márgenes estándar y tipografía moderna."""
        doc = docx.Document()
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
        return doc

    def format_run(self, run, font_name="Calibri", size_pt=11, bold=False, color_rgb=(0, 0, 0)):
        """Aplica formato visual a un fragmento de texto (run)."""
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        run.bold = bold
        run.font.color.rgb = RGBColor(*color_rgb)

    def add_title(self, doc: docx.Document, title_text: str):
        """Agrega título principal del documento."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title_text)
        self.format_run(run, font_name="Calibri", size_pt=16, bold=True, color_rgb=(30, 41, 59))
        p.paragraph_format.space_after = Pt(14)

    def add_printed_text(self, doc: docx.Document, text: str):
        """Agrega un párrafo de texto impreso estático."""
        p = doc.add_paragraph()
        run = p.add_run(text)
        self.format_run(run, font_name="Calibri", size_pt=11, color_rgb=(51, 65, 85))
        p.paragraph_format.space_after = Pt(6)

    def add_editable_field(self, doc: docx.Document, field: Dict[str, Any]):
        """
        Inserta un campo editable según su tipo:
        - texto: "Etiqueta: [____________________]"
        - fecha: "Fecha: [ DD / MM / AAAA ]"
        - firma: "Firma: _________________________"
        - casilla: "[  ] Opción / Etiqueta"
        """
        tipo = field.get("tipo_campo", "texto").lower()
        etiqueta = field.get("etiqueta", "Campo").strip()
        if not etiqueta.endswith(":") and tipo != "casilla":
            etiqueta += ":"

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)

        if tipo == "casilla":
            run_box = p.add_run("[  ] ")
            self.format_run(run_box, font_name="Calibri", size_pt=12, bold=True, color_rgb=(15, 23, 42))
            run_label = p.add_run(etiqueta)
            self.format_run(run_label, font_name="Calibri", size_pt=11, color_rgb=(30, 41, 59))

        elif tipo == "fecha":
            run_lbl = p.add_run(f"{etiqueta} ")
            self.format_run(run_lbl, font_name="Calibri", size_pt=11, bold=True, color_rgb=(30, 41, 59))
            run_val = p.add_run("[  DD  /  MM  /  AAAA  ]")
            self.format_run(run_val, font_name="Calibri", size_pt=11, color_rgb=(100, 116, 139))

        elif tipo == "firma":
            run_lbl = p.add_run(f"{etiqueta}\n\n")
            self.format_run(run_lbl, font_name="Calibri", size_pt=11, bold=True, color_rgb=(30, 41, 59))
            run_line = p.add_run("_________________________________________\n(Firma del Responsable)")
            self.format_run(run_line, font_name="Calibri", size_pt=10, color_rgb=(100, 116, 139))

        else: # texto corto o general
            run_lbl = p.add_run(f"{etiqueta} ")
            self.format_run(run_lbl, font_name="Calibri", size_pt=11, bold=True, color_rgb=(30, 41, 59))
            run_line = p.add_run("__________________________________________________")
            self.format_run(run_line, font_name="Calibri", size_pt=11, color_rgb=(148, 163, 184))

    def export_to_file(
        self,
        document_title: str,
        printed_texts: List[str],
        fields: List[Dict[str, Any]],
        output_path: str,
        images_to_embed: Optional[List[str]] = None
    ) -> str:
        """
        Construye el archivo .docx final e inserta imágenes/firmas si se proveen.
        """
        doc = self.create_base_document()

        # Título si existe
        if document_title:
            self.add_title(doc, document_title)

        # Textos impresos estáticos
        for text in printed_texts:
            if text.strip():
                self.add_printed_text(doc, text)

        # Campos editables detectados
        if fields:
            p_head = doc.add_paragraph()
            p_head.paragraph_format.space_before = Pt(12)
            run_h = p_head.add_run("CAMPOS DILIGENCIABLES DETECTADOS:")
            self.format_run(run_h, font_name="Calibri", size_pt=10, bold=True, color_rgb=(71, 85, 105))

            for field in fields:
                self.add_editable_field(doc, field)

        # Inserción de Imágenes Anexas (Logos/Firmas)
        if images_to_embed:
            p_img = doc.add_paragraph()
            p_img.paragraph_format.space_before = Pt(14)
            run_img_h = p_img.add_run("IMÁGENES Y LOGOTIPOS INCRUSTADOS:")
            self.format_run(run_img_h, font_name="Calibri", size_pt=10, bold=True, color_rgb=(71, 85, 105))

            for img_path in images_to_embed:
                if os.path.exists(img_path):
                    try:
                        doc.add_picture(img_path, width=Inches(3.0))
                    except Exception as e:
                        pass

        # Crear carpeta destino si no existe
        os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
        doc.save(output_path)
        return output_path
